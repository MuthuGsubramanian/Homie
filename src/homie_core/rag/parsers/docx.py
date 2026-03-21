from __future__ import annotations

from pathlib import Path

from homie_core.rag.parsers import ParsedDocument, TableData, TextBlock, register_parser


@register_parser("docx")
def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError:
        return ParsedDocument(source_path=str(path), metadata={"error": "python-docx not installed"})
    doc = Document(str(path))
    blocks = []
    tables = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style:
            level = int(style.replace("Heading", "").strip() or "1")
            blocks.append(TextBlock(content=para.text, block_type="heading", level=level))
        else:
            blocks.append(TextBlock(content=para.text, block_type="paragraph"))
    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
        rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
        tables.append(TableData(headers=headers, rows=rows))
    metadata = {"format": "docx"}
    try:
        props = doc.core_properties
        metadata["title"] = props.title or ""
        metadata["author"] = props.author or ""
    except Exception:
        pass
    return ParsedDocument(text_blocks=blocks, metadata=metadata, tables=tables, source_path=str(path))
