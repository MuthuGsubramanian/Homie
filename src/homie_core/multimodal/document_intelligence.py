from __future__ import annotations

import logging
import os
from typing import Any, Callable, Optional

logger = logging.getLogger(__name__)

try:
    import csv as _csv
except ImportError:  # pragma: no cover — stdlib, always present
    pass

try:
    from docx import Document as _DocxDocument

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import pdfplumber as _pdfplumber

    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

# Keyword sets used by the lightweight classifier
_DOC_TYPE_KEYWORDS: dict[str, list[str]] = {
    "invoice": ["invoice", "bill to", "amount due", "total due", "payment terms"],
    "receipt": ["receipt", "paid", "transaction", "thank you for your purchase"],
    "contract": ["agreement", "contract", "parties", "whereas", "terms and conditions", "hereby"],
    "letter": ["dear", "sincerely", "regards", "to whom it may concern"],
    "report": ["report", "findings", "analysis", "conclusion", "executive summary"],
    "statement": ["statement", "balance", "account", "period ending"],
}


class DocumentIntelligence:
    """Deep document understanding for financial, legal, and business docs."""

    def __init__(self, inference_fn: Optional[Callable[..., str]] = None):
        self._infer = inference_fn

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract_tables(self, doc_path: str) -> list[list[list[str]]]:
        """Extract tables from a PDF or DOCX.

        Returns a list of tables, where each table is a list of rows, and each
        row is a list of cell strings.
        """
        self._ensure_file(doc_path)
        ext = os.path.splitext(doc_path)[1].lower()

        if ext == ".pdf":
            return self._tables_from_pdf(doc_path)
        if ext in (".docx", ".doc"):
            return self._tables_from_docx(doc_path)

        logger.debug("Unsupported file type for table extraction: %s", ext)
        return []

    def extract_key_values(self, doc_path: str) -> dict[str, str]:
        """Extract key-value pairs from forms / invoices.

        Uses the LLM when available; otherwise falls back to simple colon-
        delimited line parsing on the extracted text.
        """
        self._ensure_file(doc_path)
        text = self._extract_text(doc_path)

        if self._infer is not None:
            try:
                prompt = (
                    "Extract all key-value pairs from this document text. "
                    "Return them as a JSON object.\n\n" + text[:4000]
                )
                import json

                raw = self._infer(prompt)
                # Try to parse the LLM output as JSON
                return json.loads(raw)
            except Exception:
                logger.debug("LLM key-value extraction failed, using fallback", exc_info=True)

        return self._heuristic_key_values(text)

    def classify_document(self, doc_path: str) -> str:
        """Classify a document as invoice, contract, letter, report, statement, or receipt."""
        self._ensure_file(doc_path)
        text = self._extract_text(doc_path).lower()

        if self._infer is not None:
            try:
                prompt = (
                    "Classify this document into exactly one of these categories: "
                    "invoice, contract, letter, report, statement, receipt.\n\n"
                    + text[:3000]
                )
                result = self._infer(prompt).strip().lower()
                for dtype in _DOC_TYPE_KEYWORDS:
                    if dtype in result:
                        return dtype
            except Exception:
                logger.debug("LLM classification failed, using fallback", exc_info=True)

        return self._heuristic_classify(text)

    def summarize_document(self, doc_path: str) -> str:
        """Generate an intelligent summary of the document."""
        self._ensure_file(doc_path)
        text = self._extract_text(doc_path)

        if self._infer is not None:
            try:
                prompt = (
                    "Provide a concise summary of this document in 2-4 sentences.\n\n"
                    + text[:5000]
                )
                return self._infer(prompt)
            except Exception:
                logger.debug("LLM summary failed, using fallback", exc_info=True)

        # Fallback — return first meaningful lines
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        preview = " ".join(lines[:10])
        if len(preview) > 500:
            preview = preview[:500] + "..."
        return preview if preview else "[Empty document]"

    # ------------------------------------------------------------------
    # Text extraction helpers
    # ------------------------------------------------------------------

    def _extract_text(self, doc_path: str) -> str:
        ext = os.path.splitext(doc_path)[1].lower()
        if ext == ".pdf":
            return self._text_from_pdf(doc_path)
        if ext in (".docx", ".doc"):
            return self._text_from_docx(doc_path)
        if ext == ".txt":
            with open(doc_path, encoding="utf-8", errors="replace") as f:
                return f.read()
        if ext == ".csv":
            with open(doc_path, encoding="utf-8-sig") as f:
                return f.read()
        logger.debug("Unsupported extension for text extraction: %s", ext)
        return ""

    @staticmethod
    def _text_from_pdf(path: str) -> str:
        if not _HAS_PDFPLUMBER:
            logger.debug("pdfplumber not installed — cannot extract PDF text")
            return ""
        try:
            pages: list[str] = []
            with _pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)
        except Exception:
            logger.debug("PDF text extraction failed", exc_info=True)
            return ""

    @staticmethod
    def _text_from_docx(path: str) -> str:
        if not _HAS_DOCX:
            logger.debug("python-docx not installed — cannot extract DOCX text")
            return ""
        try:
            doc = _DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            logger.debug("DOCX text extraction failed", exc_info=True)
            return ""

    # ------------------------------------------------------------------
    # Table extraction helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _tables_from_pdf(path: str) -> list[list[list[str]]]:
        if not _HAS_PDFPLUMBER:
            return []
        try:
            tables: list[list[list[str]]] = []
            with _pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    for table in page.extract_tables():
                        cleaned = [[cell or "" for cell in row] for row in table]
                        tables.append(cleaned)
            return tables
        except Exception:
            logger.debug("PDF table extraction failed", exc_info=True)
            return []

    @staticmethod
    def _tables_from_docx(path: str) -> list[list[list[str]]]:
        if not _HAS_DOCX:
            return []
        try:
            doc = _DocxDocument(path)
            tables: list[list[list[str]]] = []
            for tbl in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in tbl.rows]
                tables.append(rows)
            return tables
        except Exception:
            logger.debug("DOCX table extraction failed", exc_info=True)
            return []

    # ------------------------------------------------------------------
    # Heuristic helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _heuristic_key_values(text: str) -> dict[str, str]:
        """Parse 'Key: Value' lines."""
        pairs: dict[str, str] = {}
        for line in text.splitlines():
            if ":" in line:
                parts = line.split(":", 1)
                key = parts[0].strip()
                value = parts[1].strip()
                if key and value and len(key) < 60:
                    pairs[key] = value
        return pairs

    @staticmethod
    def _heuristic_classify(text: str) -> str:
        """Score each document type by keyword hits."""
        best_type = "report"
        best_score = 0
        for doc_type, keywords in _DOC_TYPE_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in text)
            if score > best_score:
                best_score = score
                best_type = doc_type
        return best_type

    @staticmethod
    def _ensure_file(path: str) -> None:
        if not os.path.isfile(path):
            raise FileNotFoundError(f"Document not found: {path}")
